from docx import Document
from pathlib import Path

p = Path(r'c:\Users\TIAN\Desktop\AI\合作开发协议.docx')
doc = Document(p)

full = []
for i, para in enumerate(doc.paragraphs):
    txt = para.text.strip()
    if txt:
        full.append(f"[{i}] {txt}")

tbl_texts = []
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            t = cell.text.strip()
            if t:
                tbl_texts.append(f"[T{ti} R{ri} C{ci}] {t}")

print("===== 段落 =====")
print("\n".join(full))
print("\n===== 表格 =====")
print("\n".join(tbl_texts))
