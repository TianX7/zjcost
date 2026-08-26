from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re

ROOT = Path(r'c:\Users\TIAN\Desktop\AI\AI\zjcost-main\zjcost-main')
OUT_DIR = Path(r'c:\Users\TIAN\Desktop\AI\软著材料')
OUT_DIR.mkdir(exist_ok=True)

CODE_FILES = [
    ROOT / 'backend' / 'app' / 'main.py',
    ROOT / 'backend' / 'app' / 'db' / 'base.py',
    ROOT / 'backend' / 'app' / 'services' / 'pricing_engine.py',
    ROOT / 'backend' / 'app' / 'api' / 'routes' / 'projects.py',
    ROOT / 'backend' / 'app' / 'smart' / 'framework' / 'base_agent.py',
    ROOT / 'backend' / 'app' / 'smart' / 'agents' / 'v2' / 'orchestrator.py',
    ROOT / 'frontend' / 'src' / 'App.tsx',
    ROOT / 'frontend' / 'src' / 'pages' / 'Dashboard.tsx',
    ROOT / 'frontend' / 'src' / 'pages' / 'SystemSettings.tsx',
    ROOT / 'frontend' / 'src' / 'components' / 'Ifc3DViewer.tsx',
    ROOT / 'frontend' / 'src' / 'api.ts',
]


def set_cell_font(cell, font_name='宋体', size=10.5):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(size)


def set_cell_no_padding(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ['top', 'left', 'bottom', 'right']:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), '0')
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def add_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9)

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_header(section, text='筑衡 V1.0 · 操作手册'):
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hrun = hp.add_run(text)
    hrun.font.name = '宋体'
    hrun._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    hrun.font.size = Pt(9)


def add_toc(doc):
    """插入 Word 目录字段（转换 PDF 时由 Word 自动填充）。"""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def generate_code_pdf():
    """使用 reportlab 直接生成程序鉴别材料 PDF。

    - 每页 50 行代码
    - 页眉显示软件名称"筑衡 V1.0"
    - 页脚显示页码
    - 前 30 页 + 后 30 页（不足 60 页则全部输出）
    - 代码使用等宽字体 Consolas，中文使用宋体
    """
    print("正在生成程序鉴别材料...")
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # 注册字体：Consolas（等宽，用于代码）+ SimSun（宋体，用于中文）
    FONT_CODE = 'Consolas'
    FONT_CJK = 'SimSun'
    fonts_dir = Path(os.environ.get('WINDIR', r'C:\Windows')) / 'Fonts'
    pdfmetrics.registerFont(TTFont(FONT_CODE, str(fonts_dir / 'consola.ttf')))
    pdfmetrics.registerFont(TTFont(FONT_CJK, str(fonts_dir / 'simsun.ttc'), subfontIndex=0))

    # 页面布局常量（A4: 595.27 x 841.89 pt）
    PAGE_W, PAGE_H = A4
    MARGIN_L = 40
    MARGIN_R = 40
    FONT_SIZE = 8
    LINE_HEIGHT = 13.5
    LINES_PER_PAGE = 50
    HEADER_TEXT = '筑衡 V1.0'

    # 代码起始 y 坐标（页眉下方）
    code_top_y = PAGE_H - 55
    # 代码区域底部（页脚上方留白）
    code_bottom_y = 40
    # 校验：50 行是否能排下
    needed = (LINES_PER_PAGE - 1) * LINE_HEIGHT + FONT_SIZE
    available = code_top_y - code_bottom_y
    if needed > available:
        print(f"  警告: 每页 {LINES_PER_PAGE} 行需要 {needed:.0f}pt，可用 {available:.0f}pt，将缩减行高")
        LINE_HEIGHT = (available - FONT_SIZE) / (LINES_PER_PAGE - 1)

    # 收集所有源码行
    lines = []
    for f in CODE_FILES:
        if not f.exists():
            print(f"  跳过不存在文件: {f}")
            continue
        text = f.read_text(encoding='utf-8', errors='ignore')
        rel = f.relative_to(ROOT)
        lines.append(f"# === FILE: {rel} ===")
        lines.extend(text.splitlines())
        lines.append("")

    total = len(lines)
    print(f"  源码总行数: {total}")

    # 前 30 页 + 后 30 页
    page_capacity = LINES_PER_PAGE * 30
    if total <= page_capacity:
        selected = lines
    else:
        front = lines[:page_capacity]
        back = lines[-page_capacity:]
        selected = front + back

    total_pages = (len(selected) + LINES_PER_PAGE - 1) // LINES_PER_PAGE
    print(f"  计划总页数: {total_pages}（每页 {LINES_PER_PAGE} 行）")

    # CJK 字符检测（含中文、全角标点、CJK 标点）
    cjk_re = re.compile(r'[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef\u2018-\u201f]')

    def draw_mixed(c, x, y, text, size):
        """绘制可能包含中英文混排的代码行：ASCII 用 Consolas，CJK 用宋体。"""
        if not text:
            return
        cur_x = x
        buf = []
        buf_is_cjk = None
        for ch in text:
            is_cjk = bool(cjk_re.match(ch))
            if buf_is_cjk is None:
                buf_is_cjk = is_cjk
                buf.append(ch)
            elif is_cjk == buf_is_cjk:
                buf.append(ch)
            else:
                seg = ''.join(buf)
                font = FONT_CJK if buf_is_cjk else FONT_CODE
                c.setFont(font, size)
                c.drawString(cur_x, y, seg)
                cur_x += c.stringWidth(seg, font, size)
                buf = [ch]
                buf_is_cjk = is_cjk
        if buf:
            seg = ''.join(buf)
            font = FONT_CJK if buf_is_cjk else FONT_CODE
            c.setFont(font, size)
            c.drawString(cur_x, y, seg)

    out_path = OUT_DIR / '程序鉴别材料.pdf'
    c = canvas.Canvas(str(out_path), pagesize=A4)

    for page_idx in range(total_pages):
        page_lines = selected[page_idx * LINES_PER_PAGE:(page_idx + 1) * LINES_PER_PAGE]

        # 页眉：左侧软件名称，右侧页码信息
        c.setFont(FONT_CJK, 9)
        c.drawString(MARGIN_L, PAGE_H - 30, HEADER_TEXT)
        c.setFont(FONT_CODE, 8)
        c.drawRightString(PAGE_W - MARGIN_R, PAGE_H - 30, f'Page {page_idx + 1} / {total_pages}')
        # 页眉下划线
        c.setStrokeColorRGB(0.4, 0.4, 0.4)
        c.setLineWidth(0.5)
        c.line(MARGIN_L, PAGE_H - 36, PAGE_W - MARGIN_R, PAGE_H - 36)

        # 代码行
        y = code_top_y
        for line in page_lines:
            draw_mixed(c, MARGIN_L, y, line.replace('\t', '    '), FONT_SIZE)
            y -= LINE_HEIGHT

        # 页脚页码
        c.setFont(FONT_CJK, 8)
        c.drawCentredString(PAGE_W / 2, 22, f'- {page_idx + 1} -')

        c.showPage()

    c.save()
    print(f"  已保存: {out_path}")
    print(f"  实际总页数: {total_pages}")
    return out_path


def generate_manual_pdf():
    print("正在生成操作手册...")
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)
    # 首页（封面）不显示页眉页脚
    section.different_first_page_header_footer = True
    add_header(section)
    add_page_number(section)

    def add_heading(text, level=1):
        p = doc.add_heading(text, level=level)
        for r in p.runs:
            r.font.name = '黑体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        return p

    def add_para(text, bold=False, indent=True, size=12):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(size)
        run.font.bold = bold
        if indent:
            p.paragraph_format.first_line_indent = Inches(0.3)
        p.paragraph_format.line_spacing = 1.5
        return p

    # ===== 封面页 =====
    # 顶部留白
    for _ in range(4):
        doc.add_paragraph()
    # 软件名称
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('筑衡 V1.0')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(36)
    run.font.bold = True
    # 文档名称
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Inches(0.3)
    run = p.add_run('操作手册')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(28)
    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Inches(0.2)
    run = p.add_run('全过程工程造价协同管控平台')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(14)
    # 底部日期
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2026 年 06 月')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(14)
    doc.add_page_break()

    # ===== 目录页 =====
    add_heading('目录', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_toc(doc)
    doc.add_page_break()

    chapters = [
        ('第一章 系统概述', [
            '筑衡：全过程工程造价协同管控平台，面向建设工程造价管理领域，提供从图纸识别、清单计价、审计复核到三维漫游核查的全流程数字化能力。',
            '系统采用 B/S 架构，前端基于 React + TypeScript + Three.js 构建，后端基于 Python FastAPI + SQLAlchemy + SQLite 构建，支持跨平台访问。',
            '平台核心目标是通过图纸智能识别、智能定额匹配、造价风险审计和 BIM 可视化，提升造价编制与审核效率，降低人工算量误差。',
            '系统主要面向建筑工程造价管理、工程建设信息化领域，适用于建设项目投资决策、设计、招投标、施工及竣工结算等全生命周期造价管理场景。',
            '平台支持多用户协同工作，提供项目台账管理、图纸智能识别、清单计价、审计复核、成果报表、IFC 三维漫游核查等核心功能模块。',
        ]),
        ('第二章 系统架构与技术栈', [
            '系统采用前后端分离架构，前端使用 React 18 + TypeScript + Ant Design + Vite 构建，使用 CSS 样式表进行样式管理，使用 Three.js 实现 BIM 三维可视化。',
            '后端使用 Python 3.12 + FastAPI + SQLAlchemy + Pydantic 构建，默认使用 SQLite 数据库，通过 Alembic 进行数据库迁移管理。',
            '智能层采用自研多 Agent 框架，包含 Orchestrator 意图路由、12 个以上专业 Agent（估价、校验、分析、批量审核等）、30 个以上工具集，支持三层模型路由（fast/balanced/powerful）。',
            '系统支持跨平台部署，可在 Windows 10/11 和 Linux Ubuntu 22.04 上运行，通过 Chrome、Edge、Firefox 等现代浏览器访问。',
            '后端支持通过 Nginx 反向代理部署，支持速率限制、CORS 跨域配置、JWT 身份认证等安全机制。',
        ]),
        ('第三章 登录与首页', [
            '用户通过浏览器访问系统地址，进入登录页面。输入账号密码完成登录后，系统跳转至造价总控台首页。',
            '造价总控台首页展示项目总数、造价总额、待审事项、风险预警等关键 KPI 指标，KPI 卡片支持数字滚动动效展示。',
            '左侧导航栏按功能分组，提供造价总览（造价总控台、项目台账）、工程与数据（算量中心、数据资源）、计价与复核（计价与复核）、系统管理（平台配置）等功能入口。',
            '首页底部流程条可快速跳转到下一步操作页面，引导用户按流程完成造价编制工作。',
            '系统支持演示模式，开启后自动加载演示数据快照，隐藏调试元素，放大关键信息展示，适用于比赛现场和汇报演示场景。',
        ]),
        ('第四章 项目台账管理', [
            '项目台账模块提供项目全生命周期管理能力。用户可创建新项目，填写项目名称、建设地点、工程类别、建筑面积等基本信息。',
            '项目列表支持按名称搜索、按状态筛选、按创建时间排序，卡片式展示项目概况和进度。',
            '点击项目卡片进入项目详情页，可查看项目基本信息、清单条目、定额绑定、计算结果、审计记录等完整信息。',
            '支持项目成员管理，可添加、移除项目成员并分配角色权限。',
            '支持项目快照功能，可保存项目当前状态，便于版本对比和回溯。',
        ]),
        ('第五章 图纸智能识别', [
            '进入图纸识别模块，系统默认展示建筑专业 CAD 预览底图。顶部工具条支持切换建筑、给排水、电气、暖通、消防五个专业。',
            '点击“选择图纸文件”上传 DWG、DXF、PDF 或 PNG 格式图纸，系统自动完成构件识别、清单生成和定额匹配。',
            '识别引擎通过 DWG/DXF 解析、OCR 文字提取与图例解析，自动提取柱、梁、墙、板等构件信息，生成工程量清单。',
            '识别完成后右侧解析状态面板显示识别进度、构件数量、清单条目和造价概览，支持展开查看明细。',
            '用户可对识别结果进行手动调整，包括修正构件属性、新增/删除/合并构件，确保识别结果准确。',
            '识别结果可一键导出为工程量清单，进入清单计价模块进行后续计价工作。',
        ]),
        ('第六章 算量中心', [
            '算量中心模块集中管理项目的所有工程量数据，支持按专业、按楼层、按构件分类查看。',
            '支持从图纸识别结果一键导入工程量，也支持手动录入或通过 Excel 批量导入。',
            '提供工程量校验功能，自动检查异常工程量（如负值、过大/过小值），并给出预警。',
            '支持工程量快照对比，可对比不同版本的工程量变化。',
        ]),
        ('第七章 清单计价', [
            '清单计价模块支持从图纸识别结果一键导入工程量，按专业分类展示分部分项工程量清单。',
            '系统通过智能定额匹配引擎，为每条清单项推荐最佳定额，支持人工调整和系数设置。',
            '用户可对清单条目进行单价调整、定额替换、综合单价分析和费用汇总。系统实时计算合价并生成造价汇总表。',
            '综合单价分析展示人工费、材料费、机械费、管理费、利润、规费、税金等完整费用分解。',
            '支持市场价动态查询，可从广材网、我的钢铁网、造价通等数据源获取最新材料价格。',
            '支持导出 Excel/PDF 报表，便于线下复核与归档。',
        ]),
        ('第八章 审计复核', [
            '审计复核模块对造价结果进行智能校验，识别清单缺项、单价异常、工程量偏差、重复列项等风险点。',
            '系统按问题等级（高、中、低）生成审计报告，列出问题描述、涉及清单、建议处理措施和责任人。',
            '审计人员可逐项确认、驳回或修正，最终生成审计结论。',
            '支持审计轨迹追踪，记录所有审计操作的时间、人员、内容，确保审计过程可追溯。',
            '审计报告支持导出 PDF 格式，便于归档和汇报。',
        ]),
        ('第九章 成果报表', [
            '成果报表模块汇总项目全周期造价数据，生成项目总造价表、专业造价分析表、费用构成表、对比分析表等多维度报表。',
            '支持按时间范围、专业、楼栋等维度筛选，支持导出 PDF 和 Excel 格式。',
            '报表中心提供项目总览、绑定进度、异常报告、历史对比等可视化图表。',
            '支持自定义报表模板，满足不同项目的报表需求。',
        ]),
        ('第十章 IFC 三维漫游核查', [
            '进入 IFC 三维漫游模块，系统加载建筑 IFC 模型，提供第一人称和第三人称两种漫游视角。',
            '用户可使用 WASD 键移动，鼠标控制视角，点击构件查看属性信息。系统支持自动巡航模式，按预设视点飞行展示关键空间。',
            '漫游过程中会自动识别并标注模型中的问题构件，辅助现场核查与质量问题定位。',
            '支持模型层级树浏览，可按楼层、按专业筛选构件显示。',
            '支持测量工具，可测量构件尺寸、距离和面积。',
        ]),
        ('第十一章 数据资源管理', [
            '数据资源模块集中管理定额库、材料价格库、清单标准等基础数据。',
            '定额库支持按专业分类浏览，支持导入 2020 土建定额、安装定额等标准定额数据。',
            '材料价格库支持手动维护和自动采集，记录价格历史趋势。',
            '清单标准编码库支持国标清单编码管理，确保清单编制规范统一。',
            '支持数据备份与恢复，保障项目数据安全。',
        ]),
        ('第十二章 平台配置', [
            '管理员可在平台配置模块管理用户权限、专业分类、定额库、报表模板等基础数据。',
            '系统设置支持配置费率规则（管理费、利润、税金）以及推理服务 A/B/C/D 的 API Key、Base URL、模型名称等参数，可在推理服务 A、推理服务 B、推理服务 C、推理服务 D 及兼容模式之间切换。',
            '支持用户角色管理，可分配管理员、造价员、审计员等不同角色权限。',
            '系统支持环境变量配置，包括数据库连接、CORS 跨域、认证开关、自动建表等选项。',
        ]),
        ('第十三章 演示模式', [
            '系统提供演示模式开关，开启后自动加载演示数据快照，隐藏调试元素，放大关键信息展示。',
            '演示导览面板按步骤引导讲解，支持计时器、快捷键导航和流程串联，适用于比赛现场和汇报演示场景。',
            '演示模式提供数字展厅功能，可展示项目全流程操作步骤和关键成果。',
            '支持演示流程自定义，可按需调整演示步骤和顺序。',
        ]),
        ('第十四章 常见问题', [
            '若图纸上传后解析失败，请确认文件格式是否正确（支持 DWG、DXF、PDF、PNG），或尝试转换为 PDF 后重新上传。',
            '若三维漫游加载缓慢，建议检查显卡是否支持 WebGL，或降低模型精度后重试。',
            '若定额匹配结果不准确，可手动调整定额绑定，或更新定额库数据后重新匹配。',
            '若计算结果异常，请检查清单项的工程量、定额绑定和费率设置是否正确。',
            '若系统登录失败，请检查账号密码是否正确，或联系管理员重置密码。',
            '更多问题请联系系统管理员或技术支持团队。',
        ]),
    ]

    for title, paras in chapters:
        add_heading(title, level=2)
        for para in paras:
            add_para(para)

    # 附录
    add_heading('附录 A 术语表', level=2)
    terms = [
        'BIM：建筑信息模型，是以建筑工程项目的各项相关信息数据为基础，通过数字信息仿真模拟建筑物所具有的真实信息。',
        'IFC：Industry Foundation Classes，国际通用的 BIM 数据交换标准。',
        'DWG：AutoCAD 的原生图形文件格式。',
        'DXF：Drawing Exchange Format，AutoCAD 的图形交换文件格式。',
        '清单计价：按照工程量清单项目设置、工程量计算规则和综合单价法计算工程造价的方法。',
        '定额：规定完成单位合格产品所需人工、材料、机械台班消耗量的标准。',
        '综合单价：完成一个规定计量单位的分部分项工程量清单项目所需的人工费、材料费、机械费、管理费和利润，并考虑风险费用。',
        'BOQ：Bill of Quantities，工程量清单。',
        'FastAPI：基于 Python 的高性能 Web 框架。',
        'React：用于构建用户界面的 JavaScript 库。',
        'Three.js：基于 WebGL 的 JavaScript 3D 图形库。',
        'SQLAlchemy：Python 的 SQL 工具包和对象关系映射（ORM）框架。',
    ]
    for t in terms:
        add_para(t)

    add_heading('附录 B 快捷键说明', level=2)
    shortcuts = [
        '图纸识别：滚轮缩放，按住鼠标拖拽平移。',
        '三维漫游：WASD 移动，Shift 加速，Esc 退出漫游，Space 切换巡航。',
        '演示导览：← 上一步，→ 下一步，Space 暂停/继续，R 重置计时。',
        '全局：Ctrl+S 保存当前操作，Ctrl+E 导出报表，Ctrl+F 搜索清单项。',
    ]
    for s in shortcuts:
        add_para(s)

    add_heading('附录 C 运行环境要求', level=2)
    env_reqs = [
        '开发硬件环境：CPU i5 及以上，内存 16GB 及以上，硬盘 512GB 及以上。',
        '运行硬件环境：CPU i3 及以上，内存 8GB 及以上，硬盘 256GB 及以上。',
        '开发操作系统：Windows 10/11 64 位、Linux Ubuntu 22.04。',
        '运行操作系统：Windows 10/11、Linux。',
        '开发环境/工具：Python 3.12+、FastAPI、SQLite、Node.js 18+、React 18+、Vite。',
        '运行支撑环境：Chrome 90+、Edge 90+、Firefox 90+、SQLite 3.35+、Nginx。',
        '编程语言：Python、JavaScript、HTML、TypeScript。',
    ]
    for r in env_reqs:
        add_para(r)

    out_path = OUT_DIR / '文档鉴别材料_操作手册.docx'
    doc.save(out_path)
    print(f"  已保存: {out_path}")


def convert_docx_to_pdf(docx_path: Path) -> Path:
    pdf_path = docx_path.with_suffix('.pdf')
    try:
        import win32com.client as wc
        word = wc.Dispatch('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        # 更新目录等字段，确保 PDF 中目录已填充
        try:
            doc.Fields.Update()
        except Exception:
            pass
        try:
            for toc in doc.TablesOfContents:
                toc.Update()
        except Exception:
            pass
        doc.SaveAs(str(pdf_path.resolve()), FileFormat=17)
        doc.Close()
        word.Quit()
    except Exception as e:
        print(f'  PDF 转换失败（请确认已安装 Word）: {e}')
    return pdf_path


if __name__ == '__main__':
    # 仅生成程序鉴别材料 PDF（reportlab 直接输出，无需 Word）
    generate_code_pdf()
    print("程序鉴别材料生成完成")
