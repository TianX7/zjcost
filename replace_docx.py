from docx import Document
from pathlib import Path

src = Path(r'c:\Users\TIAN\Desktop\AI\合作开发协议.docx')
doc = Document(src)

replacements = [
    ("智价通：智能价值评估与定价决策平台", "筑衡：全过程工程造价协同管控平台"),
    ("智价通项目", "筑衡项目"),
    ("智价通", "筑衡"),
]

def replace_text(text):
    for old, new in replacements:
        text = text.replace(old, new)
    return text

# 替换段落
for para in doc.paragraphs:
    if para.text:
        para.text = replace_text(para.text)

# 替换表格
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text:
                cell.text = replace_text(cell.text)

# 保存覆盖
doc.save(src)
print("替换完成并已保存")
